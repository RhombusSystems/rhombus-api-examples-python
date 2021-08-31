"""
    Calls to find Anomalies for Environement and Bandwidth Anomalies

"""
import os
import requests
import pandas as pd
import datetime
from contextlib import suppress
from time import time
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import IsolationForest
from docx import Document


def calc_percent_NAs(df):
    '''
    Calculates % of NAs in a all columns.
    Returns table of columns with % NA.
    '''
    nans = pd.DataFrame(df.isnull().sum().sort_values(ascending=False)/len(df), columns=['percent']) 
    idx = nans['percent'] > 0
    return nans[idx]

def convert_milli_to_date(seconds):
    '''
    Converts Milliseconds into Date Format.
    Returns Date format.
    '''
    return datetime.datetime.fromtimestamp(seconds/1000).strftime('%Y-%m-%dT%H:%M:%SZ')

def clean_date(column, date_column):
    '''
    Converts date formatting to datetime format; easier to parse date.
    Returns cleaned datetime. 
    '''
    if (column == "Temperature") or (column == "Humidity"): # Since Date formatting is different among evnironment vs. bandwidth
        clean_a = [datetime.datetime.strptime(elem, '%Y-%m-%dT%H:%M:%S') for elem in date_column] # clean up datetime format
    else: 
        clean_a = [datetime.datetime.strptime(elem, '%Y-%m-%dT%H:%M:%S %Z') for elem in date_column] # clean up datetime format
    return clean_a

def get_datetime(df):
    '''
    Reformats Time and Date column to one DATETIME column.
    Returns updates Date column.
    '''
    date, time = df["Date"], df["Time"].str.split('-').str[0]
    
    df["Date"]= date+"T"+time
    del df["Time"]

    return df["Date"]

def get_time():
    ''' 
    Finds current time and 30 days prior in milliseconds. Also converts into date format.
    Returns current millisecond, 30 days prior millisecond, current timestamp, and 30 dys prior stamp
    '''
    current_milli =  int(time() * 1000)
    thirty_days_seconds = 2592000000
    thirty_days_ago = current_milli - thirty_days_seconds
    current_milli_date = convert_milli_to_date(current_milli)
    thirty_days_ago_date = convert_milli_to_date(thirty_days_ago) 
    return current_milli, thirty_days_ago, current_milli_date, thirty_days_ago_date

def get_data(url, payload,headers,data_type):
    '''
    Grabs data via API request, creates and writes data to csv file
    Returns filename.
    '''
    response = requests.request("POST", url, json=payload, headers=headers) # Response from API data request

    current_milli, thirty_days_ago, current_milli_date, thirty_days_ago_date = get_time()
    prev_dir = os.getcwd()
    new_dir_path = prev_dir + f'/{data_type}_output'
    with suppress(FileExistsError):
        os.mkdir(f'{data_type}_output')
    os.chdir(new_dir_path)

    f_name = f'{data_type}-{thirty_days_ago_date}-to-{current_milli_date}.csv' # Filename
    f = open(f_name, "w") # Creates csv file
    f.write(response.text) # Writes data to csv file
    f.close()
    os.chdir(prev_dir)
    return f_name, new_dir_path


def wanted_anomaly_footage(perc_anomaly, column1, column2,c1_name,c2_name):
    '''
    Calculates amount of anomalies user wants and grabs X number of biggest anomalies. 
    Returns wanted anomaly data and dates.
    '''
    # Finds amount of anomaly wanted based on % specified by user
    amount_wanted_1 = round(len(column1)*(perc_anomaly/100) / 2)
    amount_wanted_2 = round(len(column2)*(perc_anomaly/100) / 2)
    
    # Sorts anomaly data columns
    c1_sorted = column1.sort_values(by=[c1_name])
    c2_sorted = column2.sort_values(by=[c2_name])
    
    # Concats most Anomalous data 
    column1_anomalies = pd.concat([c1_sorted.head(amount_wanted_1), c1_sorted.tail(amount_wanted_1)])
    column2_anomalies = pd.concat([c2_sorted.head(amount_wanted_2), c2_sorted.tail(amount_wanted_2)])
    
    # Cleans dates of most Anomalous data
    c1_date = clean_date(c1_name, column1_anomalies["Date"])
    c2_date = clean_date(c2_name, column2_anomalies["Date"])

    return column1_anomalies,c1_date, column2_anomalies, c2_date

def find_associated_camera(api_key,url,type_state):
    '''
    Finds the associated cameras to the environmental sensor.
    Returns list of camera IDs.
    '''
    headers = {
        "Accept": "application/json",
        "x-auth-scheme": "api-token",
        "Content-Type": "application/json",
        "x-auth-apikey": api_key
    }

    response = requests.request("POST", url, headers=headers)
    data = response.json()
    type_status = data[type_state] # dictionary of all data.
    camera_id_list = type_status[0].get("associatedCameras") # Grabs associatedCamera values from dictionary
    return camera_id_list

def grab_footage(api_key, device_id, duration,start_time,outlier_num,column,directory):
    '''
    Runs copy_footage_to_local_storage.py in sub-terminal and downloads footage.
    Returns None.
    '''
    # Creates path for copy_footage_to_local_storage.py to output the footage to.
    output_path = directory + f'/output{column}{outlier_num}.mp4'

    # Running copy_footage_to_local_storage.py
    os.system(f'python3 copy_footage_to_local_storage.py --api_key {api_key} --device_id {device_id} --output {output_path} --start_time {start_time} --duration {duration}')

def footage_call(column_footage_dates,api_key, device_id, duration,column, new_dir_path):

    '''
    Call to grab footage based on datetime given.
    Returns start_time; grab_footage() downloads footage to current directory.
    '''
    outlier_num = 1
    start_time = []
    directory = os.getcwd() # currrent directory
    os.chdir("..")

    # Loops through multiple footage dates wanted calling grab_footage()
    for date in column_footage_dates:
        clean_time = round(date.timestamp())
        start_time.append(clean_time)
        grab_footage(api_key, device_id, duration, clean_time,outlier_num,column,new_dir_path)

        outlier_num += 1
    os.chdir(directory)
    return start_time


def create_report_2var(graph_fname1, graph_fname2,data_type,column1_a,column2_a,new_dir_path):

    '''
    Creates report of anomalies. 
    Report contains graphs, list of user specified anomaly values, and path to footage.
    Returns None.
    '''
    os.chdir(new_dir_path)


    # Creates document and heading
    document = Document()
    document.add_heading(text=(f'{data_type} Anomaly Report'))
    
    # Adds Graphs
    document.add_paragraph('Graph of outliers')
    document.add_picture(new_dir_path+'/'+graph_fname1)
    document.add_picture(new_dir_path+'/'+graph_fname2)

    
    # Adds list of anomalies
    document.add_section()
    with pd.option_context('display.max_rows', None, 'display.max_columns', None):  # more options can be specified also
        document.add_paragraph(f'List of anomaly datetimes and values:\n{column1_a}\n{column2_a}')
    
    # Adds footage path
    document.add_section()
    document.add_paragraph(f"Footage of anomaly found at: {new_dir_path}")
    document.save(f'{data_type}AnomalyReport.docx')

def create_report_1var(graph_fname1,data_type,column1_a,new_dir_path):
    '''
    Creates report of anomalies. 
    Report contains graphs, list of user specified anomaly values, and path to footage.
    Returns None.
    '''
    os.chdir(new_dir_path)

    # Creates document and heading
    document = Document()
    document.add_heading(text=(f'{data_type} Anomaly Report'))
    
    # Adds Graphs
    document.add_paragraph('Graph of outliers')
    document.add_picture(new_dir_path+'/'+graph_fname1)
    
    # Adds list of anomalies
    document.add_section()
    with pd.option_context('display.max_rows', None, 'display.max_columns', None):  # more options can be specified also
        document.add_paragraph(f'List of anomaly datetimes and values:\n{column1_a}')
    
    # Adds footage path
    document.add_section()
    document.add_paragraph(f"Footage of anomaly found at: {new_dir_path}")

    document.save(f'{data_type}AnomalyReport.docx')

def standardize_data(data):
    '''
    Standardize data.
    Returns DataFrame of scaled data.
    '''
    np_scaled = StandardScaler().fit_transform(data)
    return pd.DataFrame(np_scaled)  

def train_forest(df,data,outliers_fraction):
    '''
    Train isolation forest.
    Returns Added column of Anomalies.
    '''
    model = IsolationForest(contamination=outliers_fraction)
    model.fit(data) 
    df['anomaly2'] = pd.Series(model.predict(data)) #finds how far point is from standardized mean
    
    return df['anomaly2']

def clean_anomaly(df,column):
    '''
    Finding and cleaning anomoly data.
    Returns cleaned data frame of anomolies.
    '''
    a = df.loc[df["anomaly2"] == -1, ['Date', column]] #anomaly
    print(f'Found {len(a)} anomalies for {column}.')
    clean_a = clean_date(column,a["Date"])

    return a, clean_a

def visualize(df, clean_dates,clean_a, a, column,output_path):

    '''
    Visualizations. Displays graph of given data and anomalies. 
    Return: None
    '''
    os.chdir(output_path)
  
    fig, ax = plt.subplots(figsize=(6,4)) # Creates plot size
    ax.plot(clean_dates, df[column], color='blue', label = 'Normal') # plots x and y 
    ax.scatter(clean_a,a[column], color='red', label = 'Anomaly') # plots anomalies

    # Format for x-axis dates
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%m/%d')) 
    _=plt.xticks(rotation=45)    

    # Labeling
    plt.xlabel('Date time')
    plt.ylabel(column)
    plt.title(f'Time Series of {column} by date time of search')

    # Display legend and plot
    plt.legend()
    plt.savefig(f"{column}_graph.jpg")

    os.chdir("..")
    return (f"{column}_graph.jpg")

def isolation_forest_test(df,data,clean_dates,column1,column2,output_path):
    '''
    Outlier test. Standardizes data, Finds Anomalies, Cleans Data and Plots graphs.
    Returns: Anomalies and Anomaly dates of columns 1 and 2.
    '''
    outliers_fraction = 0.01 # percentage of outliers estimated;; Will need adjustment based on user preference

    # Standardize Data
    std_data = standardize_data(data) 

    # Find Anomalies
    anomaly_2 = train_forest(df,std_data,outliers_fraction)

    # Cleaning anomaly data
    column1_a, column1_clean_a = clean_anomaly(df, column1)
    column2_a, column2_clean_a = clean_anomaly(df, column2) 
    
    # Plot Graph
    column1_fname = visualize(df, clean_dates, column1_clean_a, column1_a,column1,output_path)
    column2_fname = visualize(df, clean_dates, column2_clean_a, column2_a,column2,output_path)  
    
    return column1_a, column1_clean_a, column1_fname,column2_a, column2_clean_a, column2_fname

def iqr_test(df,column,data_type,output_path):
    os.chdir(output_path)
 
    q1 = column.quantile(0.25)
    q3 = column.quantile(0.75)

    iqr = q3 - q1
    iqr = iqr * 1.5

    clean_data_iqr = [] # no outliers
    diff_outliers_iqr = []   # has outliers
    diff_outlier_date = []

    count = 0
    for entry in column:
        if (entry > q3 + iqr) or (entry < q1 - iqr):
        
            date = df["Date"].values[count-1]
            diff_outlier_date.append(date)
            diff_outliers_iqr.append(entry)
        else:
            clean_data_iqr.append(entry)

        count+=1
    diff_clean_dates = [datetime.datetime.strptime(elem, '%Y-%m-%dT%H:%M:%S') for elem in df["Date"]]
    diff_outlier_dates = [datetime.datetime.strptime(elem, '%Y-%m-%dT%H:%M:%S') for elem in diff_outlier_date]

    plt.figure(figsize=(6, 8))
    plt.plot(diff_clean_dates, df['Door opened (sec)'],label='Difference Data')
    plt.xlabel('Date time')
    plt.ylabel('diff')
    _=plt.xticks(rotation=25)    
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%m/%d/%Y'))
    plt.title(f'Time Series of {data_type} Data by date time of search');
    print(f'Found {len(diff_outliers_iqr)} outliers.')
    plt.scatter(diff_outlier_dates,diff_outliers_iqr, color='red', label='Outliers')
    plt.legend(loc='upper left', frameon=True)

    plt.savefig(f"{data_type}_graph.jpg")

    os.chdir("..")
    return diff_outliers_iqr, diff_outlier_date, f"{data_type}_graph.jpg"

def seek_points(time_ms, cameraUuid, api_key):
    '''
    Creates seek point for given time and camera. Seen in console.
    Returns None: 
    '''
    url = "https://api2.rhombussystems.com/api/camera/createFootageSeekpoints"

    payload = {
        "footageSeekPoint": {
            "a": "CUSTOM",
            "ts": time_ms,
            "cdn": "Anomaly"
        },
        "cameraUuid": cameraUuid
    }
    headers = {
        "Accept": "application/json",
        "x-auth-scheme": "api-token",
        "Content-Type": "application/json",
        "x-auth-apikey": api_key
    }

    response = requests.request("POST", url, json=payload, headers=headers)
