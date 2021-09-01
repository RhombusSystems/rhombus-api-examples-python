""" 
Audit helpers.

Helper functions for audit_report and user_audit_report.
"""

import os
import json
import csv
import requests
import datetime
from time import time
from docx import Document
import matplotlib.pyplot as plt
import warnings 
warnings.filterwarnings("ignore")
from contextlib import suppress

data_type = "Audit"

def convert_milli_to_date(seconds):
    '''
    Converts Milliseconds into Date Format.
    Returns Date format.
    '''
    return datetime.datetime.fromtimestamp(seconds/1000).strftime('%Y-%m-%dT%H:%M:%SZ')

def get_time():
    ''' 
    Finds current time and 30 days prior in milliseconds. Also converts into date format.
    Returns current millisecond, 30 days prior millisecond, current timestamp, and 30 dys prior stamp
    '''
    current_milli =  int(time() * 1000)
    thirty_days_seconds = 2592000000
    thirty_days_ago = current_milli - thirty_days_seconds
    current_milli_date = convert_milli_to_date(current_milli)
    thirty_days_ago_date = convert_milli_to_date(thirty_days_ago) 
    return current_milli, thirty_days_ago, current_milli_date, thirty_days_ago_date


current_milli, thirty_days_ago, current_milli_date, thirty_days_ago_date = get_time()


def audit_grab(api_key):
    '''
    Grabs audit report for the last 30 days and writes it to a csv file.
    Returns filename to csv.
    '''
    url = "https://api2.rhombussystems.com/api/report/getAuditFeed"

    payload = {
        "timestampMsBefore": current_milli,
        "timestampMsAfter": thirty_days_ago
    }
    headers = {
        "Accept": "application/json",
        "x-auth-scheme": "api-token",
        "Content-Type": "application/json",
        "x-auth-apikey": api_key
    }

    return get_data_audit(url,payload,headers,data_type)


def get_data_audit(url,payload,headers,data_type):
    '''
    Grabs data via API request, creates and writes data to csv file
    Returns filename.
    '''
    response = requests.request("POST", url, json=payload, headers=headers)
    
    # Parsing json dictionary
    if response.status_code != 200:
        print("Encountered an Error")
    response = json.loads(response.text)
    response = response["auditEvents"]

    prev_dir = os.getcwd()
    new_dir_path = prev_dir + f'/{data_type}_output'
    with suppress(FileExistsError):
        os.mkdir(f'{data_type}_output')
    os.chdir(new_dir_path)

    # CSV file name
    f_name = f'{data_type}-{thirty_days_ago_date}-to-{current_milli_date}.csv' # Filename

    # Writing data to CSV file
    with open(f_name, 'w') as csvOutput:
        outputWriter = csv.writer(csvOutput)
        outputWriter.writerow(list(response[0].keys()))#Write Header
        for log in response[1:]:
            outputWriter.writerow(list(log.values()))#Write Data
    return f_name, new_dir_path


def clean_data_audit(df):
    '''
    Cleans Audit data. Combines city,state,country to one locations. Converts ms to date. 
    Returns smaller and cleaner audit dataframe
    '''
    df["Location"] = df["sourceCity"] +','+df["sourceState"]+','+ df["sourceCountry"]
    df['Date'] = df["timestamp"].apply(convert_milli_to_date)
    df_cleaned = df.drop(columns=['timestamp','sourceCity','sourceCountry',"sourceState","displayText","failure",'orgUuid','targetUuid','userAgent','targetName','principalType','clientType'])
    
    return df_cleaned


def find_unique_values(df,column):
    '''
    Returns unique values of wanted column.
    '''
    return df[column].unique()
    

def action_summary(df,action):
    '''
    Grabs dataframe summary of wanted action.
    Returns dataframe filtered by certain action.
    '''
    actions = find_unique_values(df,"Action")
    if action not in actions:
        return "NO such action found."
    return df.loc[df["Action"] == action]


def user_action(df,user):
    '''
    Creates a dataframe of specific user's data.
    Returns user's dataframe.
    '''
    user_df = df.loc[(df["principalName"] == user)]
    if (len(user_df) == 0):
        print("User not Found")
    return user_df


def column_activity_count(df,column):
    '''
    Returns dictionary of wanted column counts.
    '''
    return df[column].value_counts().to_dict()


def user_action_count(df,user):
    '''
    Returns actions done by a certain user including count.
    '''
    user_df = user_action(df,user)
    return column_activity_count(user_df,"action")


def user_activity_plot(user_df,user):
    xy_df = user_df[["Date","action"]]
    short = xy_df["Date"].str.split(pat="T").str[0]
    xy_df["Short Date"] = short
    date_dic = column_activity_count(xy_df,"Short Date")
    activity_count = dict(sorted(date_dic.items()))

    plt.figure(figsize=(8, 6))
    plt.plot(range(len(activity_count)), list(activity_count.values()))
    plt.xticks(range(len(activity_count)), list(activity_count.keys()))
    plt.xticks(fontsize=6)

    plt.xlabel("Dates", labelpad=20, weight='bold', size=10)
    plt.ylabel("Activity Count", labelpad=20, weight='bold', size=10)

    plt.title(f"Activity count for {user} in the past 30 Days")
    plt.savefig(f"{user}\'s_activity_graph.jpg")
    return (f"{user}\'s_activity_graph.jpg")


def user_report(user_df, user_actions, user_locations, user, graph_fname):
    '''
    Creates a report of anonymous actions, locations and dataframe. 
    '''
    # Get path.
    path = os.getcwd()

    # Creates document and heading
    document = Document()
    document.add_heading(text=(f'{user}\'s {data_type} Report'))
    

    # Adds list of Anonymous Data 
    document.add_section()
    document.add_paragraph(f'List of {user}\'s Actions:\n{user_actions}')
    
    document.add_section()
    document.add_paragraph(f'List of Locations:\n{user_locations}')

    document.add_paragraph('Graph of User Activity')
    document.add_picture(graph_fname)


    document.add_section()
    document.add_paragraph(f'All of {user}\'s Data:')
    t = document.add_table(user_df.shape[0]+1, user_df.shape[1])

    # Turn dataframe into table in document
    # add the header rows.
    for j in range(user_df.shape[-1]):
        t.cell(0,j).text = user_df.columns[j]

    # add the rest of the data frame
    for i in range(user_df.shape[0]):
        for j in range(user_df.shape[-1]):
            t.cell(i+1,j).text = str(user_df.values[i,j])
        
    
    document.save(f'{user}_{data_type}Report.docx')

    return None