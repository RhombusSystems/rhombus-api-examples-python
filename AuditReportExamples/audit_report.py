'''
    Audit Report

Program that takes in the user's api key, 
and creates an audit overview and anonymous report within the past 30 day.

Parameters: required -a API_KEY 

Downloads .docx report file, graphs, and csv of 30 days to current directory. 

Command Line Input: 
    base case: python3 audit_report.py -a {API_KEY}
'''

import os
import pandas as pd
import matplotlib.pyplot as plt
import argparse
from docx import Document
from audit_helpers import *

data_type = "Audit"

def group_users(df):
    '''
    Groups users by API, Email, Name, and Anonymous users.
    Returns all groups.
    '''
    users = find_unique_values(df,'principalName')
    api_user = []
    email_user = []
    name_user = []
    anon_user = []
    for u in users:
        if "API" in u:
            api_user.append(u)
        elif "@" in u:
            email_user.append(u)
        elif 'Anonymous' in u:
            anon_user.append(u)
        else:
            name_user.append(u)
    return api_user, email_user, name_user, anon_user

def anon_user_info(df, anon_user):
    '''
    Gets activity, locations, and summary of anonymous user.
    Returns all information.
    '''
    df_list = []
    for u in anon_user:
        df_list.append(df.loc[df["principalName"] == u])
    anon_df = pd.concat(df_list)
    anon_df_clean = anon_df.drop(columns= ["principalUuid",'principalName'])
    anon_actions = user_action_count(df,"Anonymous Share User")
    anon_locations = find_unique_values(anon_df,"Location")
    return anon_df_clean, anon_actions, anon_locations

def plot_activity(activity_count,column):
    '''
    Plots activity count by action and saves to jpg.
    Returns name of file.
    '''
    y = list(activity_count.values())
    plt.figure(figsize=(8, 6))
    plt.barh(range(len(activity_count)), list(activity_count.values()))
    plt.yticks(range(len(activity_count)), list(activity_count.keys()))
    plt.yticks(fontsize=6)

    plt.xlabel("Activity Count", labelpad=20, weight='bold', size=10)
    plt.ylabel(f"{column}", labelpad=20, weight='bold', size=10)

    locs, labels = plt.yticks()  # Get the current locations and labels.
    for i, v in enumerate(y):
        plt.text(v, i, str(v), color='red', fontsize= 'small')

    plt.title(f"Activity count per {column} in the past 30 Days")
    plt.savefig(f"{column}_graph.jpg")
    return (f"{column}_graph.jpg")

def users_in_org(api_key):
    '''
    API call to find all user emails in org.
    Returns list of user emails.
    '''
    import requests

    url = "https://api2.rhombussystems.com/api/user/getUsersInOrg"

    headers = {
    "Accept": "application/json",
    "x-auth-scheme": "api-token",
    "Content-Type": "application/json",
    "x-auth-apikey": api_key
    }

    response = requests.request("POST", url, headers=headers)
    data = response.json()
    type_status = data["users"]
    users = []
    for user in type_status:
        users.append(user.get('email'))
    return users

def inactive_user(df,api_key):
    '''
    Finds inactive users in the past 30 days.
    Returns list of inactive users.
    '''
    csv_users = find_unique_values(df,"principalName")
    org_users = users_in_org(api_key)
    inactive = []
    for user in org_users:
        if user not in csv_users:
            inactive.append(user)
    if (len(inactive) == 0):
        return None
    return inactive

def overview_report(api_user, email_user, name_user, anon_user, inactive_users, action_plot,user_plot):
    '''
    Creates a report of types of users, and plot of activity by user.
    '''
    # Get path.
    path = os.getcwd()

    # Creates document and heading
    document = Document()
    document.add_heading(text=(f'{data_type} Overview Report'))
    
    # Add Graph
    document.add_paragraph('Graph of Actions Activity')
    document.add_picture(path+'/'+action_plot)
    

    # Add Graph
    document.add_paragraph('Graph of User Activity')
    document.add_picture(path+'/'+user_plot)
    
    # Adds list of users
    document.add_section()
    document.add_paragraph(f'List of Users:\n{name_user}')
    
    document.add_section()
    document.add_paragraph(f'List of Emails:\n{email_user}')
    
    document.add_section()
    document.add_paragraph(f'List of API Users:\n{api_user}')
    
    document.add_section()
    document.add_paragraph(f'List of Anonymous Users:\n{anon_user}')
    
    document.add_section()
    document.add_paragraph(f'List of Inactive Users:\n{inactive_users}')
    
    document.save(f'{data_type}Report.docx')

    return None

def anon_report(anon_df, anon_actions, anon_locations):
    '''
    Creates a report of anonymous actions, locations and dataframe. 
    '''
    # Get path.
    path = os.getcwd()

    # Creates document and heading
    document = Document()
    document.add_heading(text=(f'{data_type} Anonymous Report'))
    

    # Adds list of Anonymous Data 
    document.add_section()
    document.add_paragraph(f'List of Anonymous Actions:\n{anon_actions}')
    
    document.add_section()
    document.add_paragraph(f'List of Locations:\n{anon_locations}')

    document.add_section()
    t = document.add_table(anon_df.shape[0]+1, anon_df.shape[1])

    # Turn dataframe into table in document
    # add the header rows.
    for j in range(anon_df.shape[-1]):
        t.cell(0,j).text = anon_df.columns[j]

    # add the rest of the data frame
    for i in range(anon_df.shape[0]):
        for j in range(anon_df.shape[-1]):
            t.cell(i+1,j).text = str(anon_df.values[i,j])
        
    document.save(f'{data_type}AnonReport.docx')

    return None

def main():

    # Parser. Gets arguements for test.
    parser = argparse.ArgumentParser(
        description='Creates a report of Audit data for the past 30 days.')
    
    parser.add_argument('--api_key', '-a', type=str, required=True,
    help='Rhombus API key')
    
    args = parser.parse_args()
    
    # Grabs Data and assigns filename
    file_name,new_dir_path = audit_grab(args.api_key)
    
    # DataFrame use for outlier test
    df = pd.read_csv(new_dir_path + '/' + file_name)


    # Clean Dataframe
    df = clean_data_audit(df)

    # Types of users
    api_user, email_user, name_user, anon_user = group_users(df)

    # Anonymous User Information
    anon_df, anon_actions, anon_locations = anon_user_info(df,anon_user)

    # Activity Count 
    activity_count = column_activity_count(df,"action")

    # User activity Count 
    user_count = column_activity_count(df,"principalName")

    # List of Inactive Users
    inactive_users = inactive_user(df,args.api_key)

    # Graph of activity per user 
    action_plot = plot_activity(activity_count,"action")
    user_plot = plot_activity(user_count,"User")

    plot_fname = user_activity_plot(anon_df,"Anonymous")

    # Write reports
    overview_report(api_user, email_user, name_user, anon_user, inactive_users, action_plot,user_plot)
    user_report(anon_df, anon_actions, anon_locations,"Anonymous",plot_fname)

if __name__ == "__main__":
    main()










